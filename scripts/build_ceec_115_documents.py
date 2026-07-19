from __future__ import annotations

import html
import json
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image

ROOT = Path(__file__).resolve().parents[1]
EXAM = ROOT / "exams" / "ceec-115-physics-g10-g11"
DATA = EXAM / "data.js"
ASSETS = EXAM / "assets" / "questions"
DOWNLOADS = EXAM / "downloads"
DELIVERY = Path.home() / "Downloads" / "ceec_115_physics_g10_g11"
SOURCE = ROOT / "source" / "ceec-115-physics"

FONT = "Microsoft JhengHei"
NAVY = "18324A"
TEAL = "0F766E"
LIGHT = "E8F3F1"
GRAY = "E9EEF2"


def load_quiz() -> dict:
    text = DATA.read_text(encoding="utf-8").strip()
    prefix = "window.QUIZ = "
    return json.loads(text[len(prefix):-1])


def set_cell_shading(cell, color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")


def set_run(run, size=10.5, bold=False, color=None) -> None:
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def configure(doc: Document, title: str) -> None:
    sec = doc.sections[0]
    sec.page_width = Cm(21.0); sec.page_height = Cm(29.7)
    sec.top_margin = Cm(1.35); sec.bottom_margin = Cm(1.35)
    sec.left_margin = Cm(1.45); sec.right_margin = Cm(1.45)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT; normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT); normal.font.size = Pt(10.5)
    for name, size, color in (("Title", 22, NAVY), ("Heading 1", 16, NAVY), ("Heading 2", 12.5, TEAL)):
        st = styles[name]; st.font.name = FONT; st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT); st.font.size = Pt(size); st.font.color.rgb = RGBColor.from_string(color)
    header = sec.header.paragraphs[0]; header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(header.add_run(title), 8, False, "607080")
    footer = sec.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE"); footer._p.append(fld)


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run(title), 22, True, NAVY)
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p2.add_run(subtitle), 11, False, TEAL)


def add_info_table(doc: Document) -> None:
    table = doc.add_table(rows=2, cols=4); table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    labels = ["班級", "____________", "姓名", "____________", "座號", "______", "得分", "____／50"]
    for i, cell in enumerate([c for row in table.rows for c in row.cells]):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER; set_cell_margins(cell)
        if i % 2 == 0: set_cell_shading(cell, LIGHT)
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(labels[i]), 10.5, i % 2 == 0, NAVY if i % 2 == 0 else None)


def add_source_note(doc: Document, quiz: dict) -> None:
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    set_run(p.add_run("篩題說明｜"), 9, True, TEAL)
    set_run(p.add_run(quiz["selectionNote"] + " 本卷保留原題號與原配分，共13題50分。"), 9, False, "52606D")


def add_image(doc: Document, name: str, max_w=6.75, max_h=8.3) -> None:
    path = ASSETS / f"{name}.jpg"
    with Image.open(path) as im:
        w, h = im.size
    width = min(max_w, max_h * w / h)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(3)
    p.add_run().add_picture(str(path), width=Inches(width))


def add_answer_line(doc: Document, label="作答") -> None:
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(5)
    set_run(p.add_run(f"{label}："), 10, True, TEAL)
    set_run(p.add_run("________________________________________________________________"), 10, False, "7B8790")


def add_writing_space(doc: Document, lines=5) -> None:
    for _ in range(lines):
        p = doc.add_paragraph("________________________________________________________________________________")
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.0
        for r in p.runs: set_run(r, 9, False, "A0A8AE")


def section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(5)
    set_run(p.add_run(text), 15, True, NAVY)


def build_exam(quiz: dict) -> Path:
    doc = Document(); configure(doc, "115分科物理｜高一高二精選卷")
    add_title(doc, "115學年度分科測驗物理科", "高一必修＋高二選修物理I、II範圍精選卷｜13題・50分")
    add_info_table(doc); add_source_note(doc, quiz)
    p=doc.add_paragraph(); set_run(p.add_run("作答提醒："),10,True,TEAL); set_run(p.add_run("單選題每題只有一個答案；多選題可複選；非選題須寫出計算或理由。"),10)

    section_heading(doc, "壹、單選題（原題1–5、10；共18分）")
    add_image(doc, "q01", max_h=2.2); add_answer_line(doc)
    add_image(doc, "g02-03-stem", max_h=1.5); add_image(doc, "g02-03-table", max_w=3.0, max_h=1.8)
    add_image(doc, "q02", max_w=5.2, max_h=3.0); add_answer_line(doc)
    add_image(doc, "q03", max_h=3.0); add_answer_line(doc)
    doc.add_page_break()
    for name in ("q04", "q05", "q10"):
        add_image(doc, name, max_h=3.4); add_answer_line(doc)

    doc.add_page_break(); section_heading(doc, "貳、多選題（原題12–14；共15分）")
    add_image(doc, "q12", max_h=4.2); add_answer_line(doc, "複選")
    add_image(doc, "q13", max_h=5.2); add_answer_line(doc, "複選")
    doc.add_page_break(); add_image(doc, "q14", max_h=4.2); add_answer_line(doc, "複選")

    section_heading(doc, "參、混合題或非選擇題（原題18–20、23；共17分）")
    add_image(doc, "g18-20-stem", max_h=5.4)
    add_image(doc, "q18", max_h=1.2); add_writing_space(doc, 4)
    add_image(doc, "q19", max_h=1.5); add_answer_line(doc)
    doc.add_page_break(); add_image(doc, "q20", max_h=3.2); add_writing_space(doc, 8)
    add_image(doc, "q23", max_h=3.0); add_writing_space(doc, 6)

    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("— 試題結束，請檢查作答 —"), 11, True, TEAL)
    path = DOWNLOADS / "ceec_115_physics_g10_g11_exam.docx"
    DOWNLOADS.mkdir(parents=True, exist_ok=True); doc.save(path); return path


def strip_html(text: str) -> str:
    text = re.sub(r"<li>", "• ", text)
    text = re.sub(r"</(?:p|li|ol|ul)>", "\n", text)
    text = re.sub(r"<[^>]+>", "", text)
    return html.unescape(re.sub(r"\n{2,}", "\n", text)).strip()


def build_answers(quiz: dict) -> Path:
    doc = Document(); configure(doc, "115分科物理｜答案與範圍審查")
    add_title(doc, "115分科物理高一高二精選卷", "答案、配分與課綱範圍審查｜v1・2026-07-19")
    add_source_note(doc, quiz)
    table = doc.add_table(rows=1, cols=6); table.style="Table Grid"; table.alignment=WD_TABLE_ALIGNMENT.CENTER
    headers=["原題","年級","課程","題型","配分","答案"]
    for c,t in zip(table.rows[0].cells,headers): set_cell_shading(c,NAVY); p=c.paragraphs[0]; set_run(p.add_run(t),9,True,"FFFFFF")
    for item in quiz["questions"]:
        cells=table.add_row().cells
        vals=[str(item["no"]),item["grade"],item["course"],{"single":"單選","multi":"多選","composite":"非選"}[item["type"]],str(item["points"]),item.get("answerText") or "見下方"]
        for c,t in zip(cells,vals): c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; set_cell_margins(c,60,70,60,70); set_run(c.paragraphs[0].add_run(t),8.3,False)
    doc.add_page_break(); section_heading(doc,"逐題解題要點")
    for item in quiz["questions"]:
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(7); p.paragraph_format.space_after=Pt(2)
        p.paragraph_format.keep_with_next = True
        set_run(p.add_run(f"原題第{item['no']}題｜{item['grade']}｜{item['points']}分｜{item.get('answerText') or '非選題'}"),11,True,NAVY)
        p2=doc.add_paragraph(); p2.paragraph_format.space_after=Pt(3)
        p2.paragraph_format.keep_with_next = True
        set_run(p2.add_run("關鍵："),9.5,True,TEAL); set_run(p2.add_run(item["keyPoint"]),9.5)
        p3=doc.add_paragraph(); p3.paragraph_format.space_after=Pt(5)
        p3.paragraph_format.keep_together = True
        set_run(p3.add_run(strip_html(item["explanationHtml"])),9.2)
    path=DOWNLOADS/"ceec_115_physics_g10_g11_answers.docx"; doc.save(path); return path


def main():
    quiz=load_quiz(); DOWNLOADS.mkdir(parents=True,exist_ok=True); DELIVERY.mkdir(parents=True,exist_ok=True)
    exam=build_exam(quiz); answers=build_answers(quiz)
    shutil.copy2(SOURCE/"ceec-115-physics-paper.pdf", DELIVERY/"official_ceec_115_physics.pdf")
    shutil.copy2(SOURCE/"ceec-115-physics-official-answers.pdf", DELIVERY/"official_ceec_115_physics_answers.pdf")
    print(json.dumps({"exam_docx":str(exam),"answers_docx":str(answers),"delivery":str(DELIVERY)},ensure_ascii=False))

if __name__=="__main__": main()
